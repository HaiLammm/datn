import json
import uuid
import logging
import re
import asyncio
from typing import Dict, Any, List
from pathlib import Path

import httpx
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import update

from app.core.config import settings
from . import models
from .rag_service import rag_service
# Story 5.4: Hybrid skill scoring integration
from .skill_scorer import skill_scorer

logger = logging.getLogger(__name__)

# Track number of pending Ollama requests for dynamic timeout calculation
# Base timeout per CV analysis - increased for long CVs with quality indicators
_BASE_TIMEOUT = 180  # seconds per CV (increased from 120 for complex analysis)
_pending_requests = 0
_request_lock = asyncio.Lock()

# Vietnamese and English section headers for CV parsing
SECTION_HEADERS_VI = [
    "thông tin cá nhân", "thông tin liên hệ", "giới thiệu bản thân",
    "mục tiêu nghề nghiệp", "mục tiêu", "tóm tắt",
    "học vấn", "trình độ học vấn", "bằng cấp",
    "kinh nghiệm làm việc", "kinh nghiệm", "lịch sử công việc",
    "kỹ năng", "kỹ năng chuyên môn", "năng lực",
    "chứng chỉ", "giấy chứng nhận", "bằng cấp chuyên môn",
    "dự án", "các dự án", "dự án tiêu biểu",
    "hoạt động", "hoạt động ngoại khóa", "hoạt động xã hội",
    "giải thưởng", "thành tích", "danh hiệu",
    "sở thích", "sở trường", "thông tin thêm",
    "người tham chiếu", "tham chiếu", "reference",
]

SECTION_HEADERS_EN = [
    "personal information", "contact information", "contact",
    "objective", "career objective", "summary", "professional summary",
    "education", "academic background", "qualifications",
    "experience", "work experience", "employment history", "professional experience",
    "skills", "technical skills", "core competencies", "competencies",
    "certifications", "certificates", "licenses",
    "projects", "key projects", "notable projects",
    "activities", "extracurricular activities", "volunteer work",
    "awards", "achievements", "honors",
    "interests", "hobbies", "additional information",
    "references", "referees",
]


class AIService:
    # Maximum characters for CV content - balanced for quality and performance
    # With smart section-based chunking, experience section is always complete
    MAX_CV_CONTENT_LENGTH = 4000  # Optimized for LLM timeout (was 5000)
    # Maximum RAG context length
    MAX_RAG_CONTEXT_LENGTH = 800  # Reduced to give more room for CV content

    def __init__(self):
        self.ollama_url = settings.OLLAMA_URL or "http://localhost:11434"
        self.model = settings.LLM_MODEL or "phi3:latest"
        self.skill_scorer = skill_scorer  # Story 5.4: Initialize hybrid skill scorer

    async def analyze_cv(
        self,
        cv_id: uuid.UUID,
        file_path: str,
        db: AsyncSession,
        force_ocr: bool = False
    ) -> None:
        """
        Analyze a CV using Ollama LLM.
        Automatically detects if OCR is needed for image-based CVs.
        Updates the database with analysis results.

        Args:
            cv_id: UUID of the CV record
            file_path: Path to the CV file
            db: Database session
            force_ocr: If True, skip text extraction and go straight to OCR
        """
        try:
            # Update status to PROCESSING
            await self._update_analysis_status(db, cv_id, models.AnalysisStatus.PROCESSING)

            cv_content = ""

            if force_ocr:
                # Directly use OCR
                logger.info(f"Force OCR mode for CV: {cv_id}")
                cv_content = await self.perform_ocr_extraction(file_path)
            else:
                # Try standard text extraction first
                try:
                    cv_content = await self._extract_text_from_file(file_path)

                    # Check if we need to fall back to OCR
                    if self.detect_if_needs_ocr(cv_content, file_path):
                        logger.info(f"Falling back to OCR for CV: {cv_id}")
                        cv_content = await self.perform_ocr_extraction(file_path)

                except Exception as e:
                    # If text extraction fails, try OCR
                    logger.warning(
                        f"Text extraction failed, trying OCR: {str(e)}")
                    cv_content = await self.perform_ocr_extraction(file_path)

            if not cv_content or len(cv_content.strip()) < 50:
                raise ValueError(
                    "Could not extract sufficient text from CV file")

            # Apply robust section splitting for better analysis
            sections = self.robust_section_split(cv_content)
            logger.info(f"Extracted {len(sections)} sections from CV")

            # Perform AI analysis
            analysis_result = await self._perform_ai_analysis(cv_content)

            # Update database with results
            await self._save_analysis_results(db, cv_id, analysis_result)

            # Update status to COMPLETED
            await self._update_analysis_status(db, cv_id, models.AnalysisStatus.COMPLETED)

        except Exception as e:
            logger.error(f"CV analysis failed for cv_id={cv_id}: {str(e)}")

            # Save fallback results so user sees something useful
            try:
                fallback_result = self._get_fallback_analysis()
                fallback_result["summary"] = f"AI analysis temporarily unavailable. Error: {str(e)[
                    :100]}"
                await self._save_analysis_results(db, cv_id, fallback_result)
                # Mark as COMPLETED with fallback data instead of FAILED
                await self._update_analysis_status(db, cv_id, models.AnalysisStatus.COMPLETED)
                logger.info(f"Saved fallback analysis for cv_id={cv_id}")
            except Exception as fallback_error:
                logger.error(f"Failed to save fallback analysis: {
                             fallback_error}")
                await self._update_analysis_status(db, cv_id, models.AnalysisStatus.FAILED)

    async def _extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text content from PDF/DOCX file.
        """
        path = Path(file_path)

        if not path.exists():
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"CV file not found: {file_path}")

        file_extension = path.suffix.lower()

        try:
            if file_extension == ".pdf":
                return await self._extract_text_from_pdf(file_path)
            elif file_extension in [".docx", ".doc"]:
                return await self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise

    async def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF

            text_parts: List[str] = []

            with fitz.open(file_path) as doc:
                for page in doc:
                    page_text = page.get_text()
                    if page_text:
                        text_parts.append(page_text)

            return "\n".join(text_parts)
        except ImportError:
            logger.error("PyMuPDF library not installed")
            raise ImportError("PyMuPDF library is required for PDF extraction")
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise

    async def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx."""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts: List[str] = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip()
                                for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)
        except ImportError:
            logger.error("python-docx library not installed")
            raise ImportError(
                "python-docx library is required for DOCX extraction")
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise

    async def _perform_ai_analysis(self, cv_content: str) -> Dict[str, Any]:
        """
        Use Ollama to analyze CV content and return structured results.
        Includes RAG context retrieval for enhanced analysis.
        Includes error handling for malformed AI responses.

        Note: Content is intelligently truncated using section-based chunking
        to preserve critical information (especially experience data).
        """
        # Step 1: Split CV into sections
        sections = self.robust_section_split(cv_content)
        logger.info(f"📑 Extracted {len(sections)} sections from CV: {
                    list(sections.keys())}")

        # Step 2: Build smart truncated content with priority sections
        priority_sections = ['experience', 'skills', 'education']
        truncated_cv = self._build_smart_truncated_content(
            sections=sections,
            priority_sections=priority_sections,
            max_length=self.MAX_CV_CONTENT_LENGTH
        )

        if len(cv_content) > self.MAX_CV_CONTENT_LENGTH:
            logger.info(f"📊 CV content intelligently truncated from {len(cv_content)} to {
                        len(truncated_cv)} chars using section-based chunking")

        # Retrieve RAG context for enhanced analysis (with length limit)
        rag_context = ""
        try:
            retrieved_docs = rag_service.retrieve_context(
                truncated_cv, top_k=2)  # Limit results
            if retrieved_docs:
                rag_context = rag_service.format_context_for_prompt(
                    retrieved_docs)
                # Truncate RAG context if too long
                if len(rag_context) > self.MAX_RAG_CONTEXT_LENGTH:
                    rag_context = rag_context[:self.MAX_RAG_CONTEXT_LENGTH] + "..."
                    logger.info(f"RAG context truncated to {
                                self.MAX_RAG_CONTEXT_LENGTH} chars")
                logger.info(f"RAG context retrieved: {
                            len(retrieved_docs)} documents")
        except Exception as e:
            logger.warning(
                f"RAG retrieval failed, proceeding without context: {str(e)}")

        # Build context section for prompt (simplified)
        context_section = ""
        if rag_context:
            context_section = f"Reference: {rag_context}\n\n"

        # Enhanced prompt that leverages section structure and extracts quality indicators
        analysis_prompt = f"""Analyze this CV with clearly marked sections and respond with ONLY a JSON object.

{context_section}IMPORTANT INSTRUCTIONS:
- The CV below is structured with section headers like [EXPERIENCE], [SKILLS], [EDUCATION]
- The [EXPERIENCE] section contains ALL job positions with date ranges
- Calculate total_years by identifying ALL date ranges (format: MM/YYYY - MM/YYYY or similar)
- Sum the duration of ALL positions to get accurate total years
- If dates span like "11/2012-08/2015" that is approximately 2.8 years
- Do NOT stop at the first position - count ALL positions in the experience section

OPTIONAL QUALITY INDICATORS (try to extract if possible):
- num_projects: count of projects mentioned
- num_awards: count of awards/achievements
- num_certifications: count of certifications
- has_leadership: true if has leadership roles (director, manager, lead, team leader, chief)
- description_quality: "good" if detailed, "medium" if average, "poor" if vague

CV CONTENT:
{truncated_cv}

Respond with JSON (quality indicators are optional):
{{"score":<0-100>,"criteria":{{"completeness":<0-100>,"experience":<0-100>,"skills":<0-100>,"professionalism":<0-100>}},"summary":"<brief>","skills":["skill1","skill2"],"experience_breakdown":{{"total_years":<sum>,"key_roles":["role1"],"industries":["ind1"],"num_projects":0,"num_awards":0,"num_certifications":0,"has_leadership":false,"description_quality":"medium"}},"strengths":["str1"],"improvements":["imp1"],"formatting_feedback":["fb1"],"ats_hints":["hint1"]}}

JSON only:"""

        try:
            result = await self._call_ollama(analysis_prompt)
            llm_analysis = self._parse_analysis_response(result)

            # Story 5.7: Apply bonus scoring for leadership and experience
            llm_analysis = self._apply_leadership_experience_bonus(llm_analysis)

            # Story 5.4: Integrate hybrid skill scoring
            try:
                skill_score_result = self.skill_scorer.calculate_skill_score(
                    cv_text=truncated_cv,
                    llm_response=llm_analysis
                )
                # Merge skill scoring results into LLM analysis
                llm_analysis["skill_breakdown"] = {
                    "completeness_score": skill_score_result["completeness_score"],
                    "categorization_score": skill_score_result["categorization_score"],
                    "evidence_score": skill_score_result["evidence_score"],
                    "market_relevance_score": skill_score_result["market_relevance_score"],
                    "total_score": skill_score_result["total_score"],
                }
                llm_analysis["skill_categories"] = skill_score_result["skill_categories"]
                llm_analysis["skill_recommendations"] = skill_score_result["recommendations"]
                logger.info(f"Hybrid skill score calculated: {
                            skill_score_result['total_score']}/25")
            except Exception as skill_err:
                logger.error(f"Skill scoring failed, using fallback: {
                             str(skill_err)}")
                # Fallback: add empty skill scoring fields
                llm_analysis["skill_breakdown"] = {
                    "completeness_score": 0,
                    "categorization_score": 0,
                    "evidence_score": 0,
                    "market_relevance_score": 0,
                    "total_score": 0,
                }
                llm_analysis["skill_categories"] = {}
                llm_analysis["skill_recommendations"] = []

            return llm_analysis
        except Exception as e:
            logger.error(f"AI analysis failed: {str(e)}")
            # Return fallback values
            return self._get_fallback_analysis()

    def _apply_leadership_experience_bonus(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        """
        Story 5.7: Apply bonus scoring for leadership roles and years of experience.

        This is a deterministic post-processing step applied after LLM analysis
        to give more weight to senior candidates.

        Bonus rules:
        - has_leadership = true: +10 to experience score (cap at 100)
        - total_years >= 5: +5 to experience score
        - total_years >= 10: +10 to experience score (replaces +5)
        - Senior titles (Director, VP, CTO, Principal): +5 to professionalism score

        Args:
            analysis: Parsed LLM analysis result.

        Returns:
            Updated analysis with adjusted scores.
        """
        try:
            criteria = analysis.get("criteria", {})
            experience_breakdown = analysis.get("experience_breakdown", {})

            experience_score = criteria.get("experience", 50)
            professionalism_score = criteria.get("professionalism", 50)

            # Track bonuses applied
            bonuses_applied = []

            # Leadership bonus
            has_leadership = experience_breakdown.get("has_leadership", False)
            if has_leadership:
                experience_score = min(100, experience_score + 10)
                bonuses_applied.append("leadership:+10")

            # Years of experience bonus
            total_years = experience_breakdown.get("total_years", 0)
            if isinstance(total_years, (int, float)):
                if total_years >= 10:
                    experience_score = min(100, experience_score + 10)
                    bonuses_applied.append("10+years:+10")
                elif total_years >= 5:
                    experience_score = min(100, experience_score + 5)
                    bonuses_applied.append("5+years:+5")

            # Senior title bonus for professionalism
            key_roles = experience_breakdown.get("key_roles", [])
            senior_titles = ["director", "vp", "vice president", "cto", "ceo",
                            "cfo", "cio", "principal", "chief", "head of"]
            has_senior_title = any(
                any(title in role.lower() for title in senior_titles)
                for role in key_roles
            )
            if has_senior_title:
                professionalism_score = min(100, professionalism_score + 5)
                bonuses_applied.append("senior_title:+5_prof")

            # Update criteria
            criteria["experience"] = experience_score
            criteria["professionalism"] = professionalism_score
            analysis["criteria"] = criteria

            # Recalculate overall score
            criteria_average = sum(criteria.values()) / len(criteria)
            analysis["score"] = int(round(criteria_average))

            if bonuses_applied:
                logger.info(f"Story 5.7 bonuses applied: {bonuses_applied}, "
                           f"new experience={experience_score}, "
                           f"new professionalism={professionalism_score}, "
                           f"new overall={analysis['score']}")

            return analysis

        except Exception as e:
            logger.warning(f"Failed to apply leadership/experience bonus: {e}")
            return analysis

    def _parse_analysis_response(self, response: str) -> Dict[str, Any]:
        """
        Parse AI response with error handling for malformed JSON.
        """
        try:
            # Try to extract JSON from response
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                data = json.loads(json_match.group())
                logger.info(
                    f"🔍 RAW LLM RESPONSE - Full JSON keys: {list(data.keys())}")
                logger.info(
                    f"🔍 RAW LLM RESPONSE - experience_breakdown: {data.get('experience_breakdown')}")
                logger.info(
                    f"🔍 RAW LLM RESPONSE - criteria: {data.get('criteria')}")
            else:
                raise ValueError("No JSON found in response")

            # Validate and sanitize the response
            # First validate experience to get quality-adjusted score
            experience_breakdown = self._validate_experience(
                data.get("experience_breakdown", {}))

            # Get initial criteria from LLM
            criteria = self._validate_criteria(data.get("criteria", {}))

            # Override experience score with quality-adjusted score
            if experience_breakdown.get("quality_adjusted_score"):
                quality_score = experience_breakdown["quality_adjusted_score"]["final_score"]
                original_score = criteria["experience"]
                criteria["experience"] = quality_score
                logger.info(f"🎯 EXPERIENCE SCORE OVERRIDE - Original: {
                            original_score} → Quality-Adjusted: {quality_score}")

            # QA Fix: Recalculate overall score from criteria average for consistency
            # This ensures the overall score reflects the actual criteria scores (including quality-adjusted experience)
            criteria_average = sum(criteria.values()) / len(criteria)
            recalculated_score = int(round(criteria_average))
            llm_original_score = self._validate_score(data.get("score", 50))

            result = {
                "score": recalculated_score,  # Use recalculated score instead of LLM's score
                "criteria": criteria,
                "summary": str(data.get("summary", "Analysis summary not available"))[:1000],
                "skills": self._validate_list(data.get("skills", []))[:50],
                "experience_breakdown": experience_breakdown,
                "strengths": self._validate_list(data.get("strengths", []))[:10],
                "improvements": self._validate_list(data.get("improvements", []))[:10],
                "formatting_feedback": self._validate_list(data.get("formatting_feedback", []))[:10],
                "ats_hints": self._validate_list(data.get("ats_hints", []))[:10],
            }

            # QA Fix: Log score consistency check
            if abs(llm_original_score - recalculated_score) > 5:
                logger.warning(
                    f"⚠️ SCORE CONSISTENCY - LLM original: {
                        llm_original_score}, "
                    f"Criteria avg: {criteria_average:.1f}, Using recalculated: {
                        recalculated_score}"
                )
            else:
                logger.info(
                    f"✓ SCORE CONSISTENT - Recalculated: {recalculated_score} (LLM: {llm_original_score})")

            logger.info(f"🔍 PARSED RESULT - Overall score: {result['score']}")
            logger.info(
                f"🔍 PARSED RESULT - Experience score (quality-adjusted): {result['criteria']['experience']}")
            logger.info(
                f"🔍 PARSED RESULT - Experience years: {result['experience_breakdown']['total_years']}")

            return result
        except (json.JSONDecodeError, ValueError) as e:
            logger.warning(f"Failed to parse AI response: {str(e)}")
            return self._get_fallback_analysis()

    def _validate_score(self, score: Any) -> int:
        """Validate score is between 0-100."""
        try:
            score_int = int(score)
            return max(0, min(100, score_int))
        except (TypeError, ValueError):
            return 50

    def _validate_criteria(self, criteria: Any) -> Dict[str, int]:
        """Validate criteria dictionary."""
        if not isinstance(criteria, dict):
            logger.warning("❌ Criteria is not a dict, using defaults")
            return {"completeness": 50, "experience": 50, "skills": 50, "professionalism": 50}

        validated = {}
        for key in ["completeness", "experience", "skills", "professionalism"]:
            raw_value = criteria.get(key, 50)
            validated_value = self._validate_score(raw_value)
            validated[key] = validated_value
            logger.info(
                f"🔍 CRITERIA DEBUG - {key}: raw={raw_value}, validated={validated_value}")

        return validated

    def _validate_list(self, items: Any) -> List[str]:
        """Validate list of strings."""
        if not isinstance(items, list):
            return []
        return [str(item)[:200] for item in items if item]

    def _validate_experience(self, exp: Any) -> Dict[str, Any]:
        """
        Validate experience breakdown and calculate quality-adjusted score.

        Extracts quality indicators (projects, awards, certifications) from LLM response
        and calculates a more accurate experience score.
        """
        if not isinstance(exp, dict):
            logger.warning(
                "❌ Experience breakdown is not a dict, using defaults")
            return {
                "total_years": 0,
                "key_roles": [],
                "industries": [],
                "num_projects": 0,
                "num_awards": 0,
                "num_certifications": 0,
                "has_leadership": False,
                "description_quality": "medium",
                "quality_adjusted_score": None,
            }

        raw_years = exp.get("total_years", 0)
        validated_years = self._validate_score(raw_years)

        logger.info(
            f"🔍 EXPERIENCE DEBUG - Raw years from LLM: {raw_years}, Validated: {validated_years}")

        # Extract quality indicators from LLM response (with safe defaults)
        try:
            num_projects = int(exp.get("num_projects", 0))
        except (ValueError, TypeError):
            num_projects = 0

        try:
            num_awards = int(exp.get("num_awards", 0))
        except (ValueError, TypeError):
            num_awards = 0

        try:
            num_certifications = int(exp.get("num_certifications", 0))
        except (ValueError, TypeError):
            num_certifications = 0

        try:
            has_leadership = bool(exp.get("has_leadership", False))
        except (ValueError, TypeError):
            has_leadership = False

        description_quality = exp.get("description_quality", "medium")

        # Validate description_quality
        if description_quality not in ["good", "medium", "poor"]:
            description_quality = "medium"

        logger.info(f"📊 QUALITY INDICATORS - Projects: {num_projects}, Awards: {num_awards}, "
                    f"Certs: {num_certifications}, Leadership: {
                        has_leadership}, "
                    f"Quality: {description_quality}")

        # Calculate quality-adjusted experience score
        quality_score_data = self._calculate_quality_adjusted_experience_score(
            total_years=validated_years,
            num_projects=num_projects,
            num_awards=num_awards,
            num_certifications=num_certifications,
            has_leadership=has_leadership,
            job_description_quality=description_quality
        )

        result = {
            "total_years": validated_years,
            "key_roles": self._validate_list(exp.get("key_roles", []))[:10],
            "industries": self._validate_list(exp.get("industries", []))[:10],
            "num_projects": num_projects,
            "num_awards": num_awards,
            "num_certifications": num_certifications,
            "has_leadership": has_leadership,
            "description_quality": description_quality,
            "quality_adjusted_score": quality_score_data,
        }

        logger.info(f"🔍 EXPERIENCE DEBUG - Full breakdown: {result}")
        return result

    def _calculate_quality_adjusted_experience_score(
        self,
        total_years: int,
        num_projects: int,
        num_awards: int,
        num_certifications: int,
        has_leadership: bool,
        job_description_quality: str = "medium"  # "poor", "medium", "good"
    ) -> Dict[str, Any]:
        """
        Calculate experience score using sigmoid/log growth model + Impact/Scope bonuses.
        
        NEW PHILOSOPHY (v2):
        - BaseScore(years): Smooth sigmoid/log curve, no sudden jumps
        - FinalScore = BaseScore + ImpactBonus + ScopeBonus
        - IC track cap at 85, Management gets separate leadership bonus
        
        Args:
            total_years: Number of years of experience
            num_projects: Number of projects mentioned (Impact indicator)
            num_awards: Number of awards/achievements (Impact indicator)
            num_certifications: Number of certifications (Scope indicator)
            has_leadership: Whether has leadership/management roles (Scope indicator)
            job_description_quality: Quality of job descriptions (Impact indicator)
            
        Returns:
            Dictionary with score, base_score, impact_bonus, scope_bonus, and explanation
        """
        import math
        
        # Step 1: Calculate smooth base score from years using hybrid sigmoid/log model
        # Target curve: 
        #   0-2y: 15-40 (junior)
        #   3-6y: 48-70 (smoothed mid, avoiding jumps)
        #   7-15y: 72-85 (senior IC cap)
        #   16+y: 85-90 (management/architect)
        
        if total_years < 1:
            base_score = 15
        elif total_years <= 2:
            # Junior ramp-up: linear 15-40
            base_score = 15 + (total_years * 12.5)
        elif total_years <= 6:
            # Smoothed mid-level growth (3-6 years): 40 → 48 → 58 → 68 → 70
            # Using piecewise linear for exact control
            milestones = {2: 40, 3: 48, 4: 58, 5: 68, 6: 70}
            if total_years in milestones:
                base_score = milestones[total_years]
            else:
                # Interpolate between milestones (shouldn't happen with int years)
                lower = int(total_years)
                upper = lower + 1
                fraction = total_years - lower
                base_score = milestones[lower] + (milestones[upper] - milestones[lower]) * fraction
        elif total_years <= 15:
            # Senior IC: logarithmic slowdown 70 → 85
            # Using log formula: 70 + 15 * log(1 + (years-6)/9)
            t_normalized = (total_years - 6) / 9.0  # Normalize 6-15 to 0-1
            base_score = 70 + 15 * math.log1p(t_normalized)  # log1p(x) = log(1+x)
        elif total_years <= 20:
            # Very senior: slow linear 85 → 87
            base_score = 85 + ((total_years - 15) * 0.4)
        else:
            # Executive: cap at 90
            base_score = 87 + min((total_years - 20) * 0.2, 3)
        
        base_score = min(90, round(base_score, 1))

        # Step 2: Calculate Impact Bonus (0-15 points)
        # Impact = Technical/product influence (projects, awards, job quality)
        impact_bonus = 0
        impact_details = []

        # Projects contribution (max 8 points) - quality over quantity
        if num_projects > 0:
            # Diminishing returns: 1→4pts, 2→6pts, 3→7pts, 4+→8pts
            projects_impact = min(num_projects * 2.5 if num_projects <= 2 else 6 + (num_projects - 2) * 0.5, 8)
            impact_bonus += projects_impact
            impact_details.append(f"Projects: +{projects_impact:.1f} ({num_projects} documented)")

        # Awards/recognition (max 5 points) - signal of exceptional impact
        if num_awards > 0:
            awards_impact = min(num_awards * 2, 5)
            impact_bonus += awards_impact
            impact_details.append(f"Awards: +{awards_impact:.1f} ({num_awards} recognitions)")

        # Job description quality (max 2 points) - depth of contribution
        quality_map = {"poor": 0, "medium": 1, "good": 2}
        desc_impact = quality_map.get(job_description_quality, 1)
        if desc_impact > 0:
            impact_bonus += desc_impact
            impact_details.append(f"Description quality: +{desc_impact} ({job_description_quality})")

        impact_bonus = min(15, impact_bonus)  # Cap at 15

        # Step 3: Calculate Scope Bonus (0-15 points)
        # Scope = Organizational influence (leadership, certifications, team scale)
        scope_bonus = 0
        scope_details = []

        # Leadership roles (max 10 points) - separate from base, scales with seniority
        if has_leadership:
            if total_years >= 15:
                leadership_scope = 10  # C-level/Director - organizational impact
            elif total_years >= 10:
                leadership_scope = 7   # Manager/Lead - team/department impact
            elif total_years >= 5:
                leadership_scope = 5   # Team Lead - small team impact
            else:
                leadership_scope = 3   # Junior lead - limited scope
            scope_bonus += leadership_scope
            scope_details.append(f"Leadership: +{leadership_scope} ({total_years}y experience)")

        # Certifications (max 5 points) - professional breadth and commitment
        if num_certifications > 0:
            # Diminishing returns: 1→2pts, 2→3pts, 3→4pts, 4+→5pts
            cert_scope = min(1 + num_certifications * 0.8, 5)
            scope_bonus += cert_scope
            scope_details.append(f"Certifications: +{cert_scope:.1f} ({num_certifications} certs)")

        scope_bonus = min(15, scope_bonus)  # Cap at 15

        # Step 4: Apply career track cap
        # IC (Individual Contributor) track: cap at 85
        # Management track (has_leadership + 10+ years): can reach 100
        is_management_track = has_leadership and total_years >= 10
        
        if is_management_track:
            # Management track: BaseScore + Impact + Scope, cap at 100
            uncapped_score = base_score + impact_bonus + scope_bonus
            final_score = min(100, uncapped_score)
            track = "Management"
        else:
            # IC track: BaseScore + Impact + Scope, cap at 85
            uncapped_score = base_score + impact_bonus + scope_bonus
            final_score = min(85, uncapped_score)
            track = "IC"
            
            # Warning if IC would exceed cap (suggests should be on management track)
            if uncapped_score > 85:
                logger.info(f"⚠️  IC candidate score capped: {uncapped_score:.0f} → 85 (consider leadership path)")

        final_score = max(0, final_score)  # Ensure non-negative

        # Step 5: Generate explanation
        explanation = f"{track} track: Base ({total_years}y)={base_score:.0f}"
        if impact_bonus > 0:
            explanation += f" + Impact={impact_bonus:.0f}"
        if scope_bonus > 0:
            explanation += f" + Scope={scope_bonus:.0f}"
        if uncapped_score != final_score:
            explanation += f" = {uncapped_score:.0f} (capped→{final_score:.0f})"
        else:
            explanation += f" = {final_score:.0f}"

        logger.info(f"📊 QUALITY-ADJUSTED SCORE - {explanation}")
        if impact_details:
            logger.info(f"   💡 Impact Bonus: {', '.join(impact_details)}")
        if scope_details:
            logger.info(f"   🎯 Scope Bonus: {', '.join(scope_details)}")

        return {
            "final_score": int(final_score),
            "base_score": int(base_score),
            "impact_bonus": int(impact_bonus),
            "scope_bonus": int(scope_bonus),
            "career_track": track,
            "explanation": explanation,
            "impact_details": impact_details,
            "scope_details": scope_details,
            # Legacy fields for backward compatibility (deprecated)
            "quality_bonus": int(impact_bonus + scope_bonus),
            "quality_penalty": 0,
            "bonus_details": impact_details + scope_details,
            "penalty_details": [],
        }

    def _get_fallback_analysis(self) -> Dict[str, Any]:
        """Return fallback analysis when AI fails."""
        return {
            "score": 50,
            "criteria": {"completeness": 50, "experience": 50, "skills": 50, "professionalism": 50},
            "summary": "Unable to generate detailed analysis. Please try again later.",
            "skills": [],
            "experience_breakdown": {"total_years": 0, "key_roles": [], "industries": []},
            "strengths": [],
            "improvements": ["Consider using a clearer format", "Add more specific details"],
            "formatting_feedback": ["Analysis service temporarily unavailable"],
            "ats_hints": ["Use standard section headers", "Include relevant keywords"],
            # Story 5.4: Include hybrid skill scoring fallback fields
            "skill_breakdown": {
                "completeness_score": 0,
                "categorization_score": 0,
                "evidence_score": 0,
                "market_relevance_score": 0,
                "total_score": 0,
            },
            "skill_categories": {},
            "skill_recommendations": [],
        }

    async def perform_ocr_extraction(self, file_path: str) -> str:
        """
        Perform OCR extraction on image-based or scanned PDF files.
        Uses EasyOCR for Vietnamese and English text recognition.

        Args:
            file_path: Path to the PDF or image file

        Returns:
            Extracted text content from OCR
        """
        path = Path(file_path)

        if not path.exists():
            logger.error(f"File not found for OCR: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            import easyocr
            from pdf2image import convert_from_path
            from PIL import Image
            import numpy as np

            # Initialize EasyOCR reader for Vietnamese and English
            reader = easyocr.Reader(['vi', 'en'], gpu=False)

            images: List[Any] = []
            file_extension = path.suffix.lower()

            if file_extension == ".pdf":
                # Convert PDF pages to images
                logger.info(f"Converting PDF to images for OCR: {file_path}")
                images = convert_from_path(file_path, dpi=300)
            elif file_extension in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
                # Direct image file
                images = [Image.open(file_path)]
            else:
                raise ValueError(f"Unsupported file format for OCR: {
                                 file_extension}")

            extracted_texts: List[str] = []

            for i, image in enumerate(images):
                logger.info(f"Processing OCR for page {i + 1}/{len(images)}")

                # Convert PIL Image to numpy array for EasyOCR
                image_np = np.array(image)

                # Perform OCR
                results = reader.readtext(image_np, detail=0, paragraph=True)

                if results:
                    page_text = "\n".join(results)
                    extracted_texts.append(page_text)

            full_text = "\n\n".join(extracted_texts)
            logger.info(f"OCR extraction completed. Total characters: {
                        len(full_text)}")

            return full_text

        except ImportError as e:
            logger.error(f"OCR library not installed: {str(e)}")
            raise ImportError(
                "OCR libraries (easyocr, pdf2image, Pillow) are required. "
                "Please install them with: pip install easyocr pdf2image Pillow"
            )
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {str(e)}")
            raise

    def detect_if_needs_ocr(self, extracted_text: str, file_path: str) -> bool:
        """
        Determine if a file needs OCR processing based on text extraction results.

        Heuristics:
        1. Text is too short (less than 100 chars)
        2. Text contains mostly garbled/unreadable characters
        3. Text has very low ratio of recognizable words
        4. PDF has images but minimal text

        Args:
            extracted_text: Text extracted via standard method
            file_path: Path to the original file

        Returns:
            True if OCR is recommended, False otherwise
        """
        # Heuristic 1: Text too short
        if len(extracted_text.strip()) < 100:
            logger.info(f"OCR needed: Text too short ({
                        len(extracted_text)} chars)")
            return True

        # Heuristic 2: Check for garbled text (high ratio of non-printable chars)
        printable_ratio = sum(1 for c in extracted_text if c.isprintable(
        ) or c.isspace()) / len(extracted_text)
        if printable_ratio < 0.8:
            logger.info(
                f"OCR needed: Low printable ratio ({printable_ratio:.2f})")
            return True

        # Heuristic 3: Check for meaningful words
        # Split into words and check if they look like real words
        words = re.findall(r'\b[a-zA-ZÀ-ỹ]{2,}\b', extracted_text)
        if len(words) < 20:
            logger.info(
                f"OCR needed: Too few recognizable words ({len(words)})")
            return True

        # Heuristic 4: Check for common CV section headers (Vietnamese or English)
        text_lower = extracted_text.lower()
        headers_found = 0
        all_headers = SECTION_HEADERS_VI + SECTION_HEADERS_EN

        for header in all_headers:
            if header in text_lower:
                headers_found += 1

        # If we found at least one section header, text extraction likely worked
        if headers_found == 0:
            logger.info(
                "OCR needed: No section headers found in extracted text")
            return True

        logger.info(f"Standard extraction sufficient: {
                    headers_found} headers found, {len(words)} words")
        return False

    def robust_section_split(self, text: str) -> Dict[str, str]:
        """
        Split CV text into sections based on Vietnamese and English headers.

        Args:
            text: Raw CV text content

        Returns:
            Dictionary mapping section names to their content
        """
        sections: Dict[str, str] = {}

        # Combine all headers and sort by length (longest first to avoid partial matches)
        all_headers = SECTION_HEADERS_VI + SECTION_HEADERS_EN
        all_headers_sorted = sorted(all_headers, key=len, reverse=True)

        # Build regex pattern to find section headers
        # Headers are typically on their own line, possibly with punctuation
        header_pattern = r'(?:^|\n)\s*(' + '|'.join(
            re.escape(h) for h in all_headers_sorted
        ) + r')[\s:：]*(?:\n|$)'

        # Find all matches
        matches = list(re.finditer(header_pattern, text,
                       re.IGNORECASE | re.MULTILINE))

        if not matches:
            # No sections found, return entire text as "content"
            return {"content": text.strip()}

        # Extract sections based on header positions
        for i, match in enumerate(matches):
            header_name = match.group(1).strip().lower()

            # Normalize header names
            header_name = self._normalize_section_header(header_name)

            # Get content between this header and the next
            start_pos = match.end()
            end_pos = matches[i + 1].start() if i + \
                1 < len(matches) else len(text)

            content = text[start_pos:end_pos].strip()

            # Store section (merge if header already exists)
            if header_name in sections:
                sections[header_name] += "\n\n" + content
            else:
                sections[header_name] = content

        return sections

    def _normalize_section_header(self, header: str) -> str:
        """
        Normalize section header to a standard name.
        Maps Vietnamese headers to their English equivalents.
        """
        header = header.lower().strip()

        # Mapping Vietnamese to English
        header_mapping = {
            "thông tin cá nhân": "personal_info",
            "thông tin liên hệ": "contact",
            "giới thiệu bản thân": "summary",
            "mục tiêu nghề nghiệp": "objective",
            "mục tiêu": "objective",
            "tóm tắt": "summary",
            "học vấn": "education",
            "trình độ học vấn": "education",
            "bằng cấp": "education",
            "kinh nghiệm làm việc": "experience",
            "kinh nghiệm": "experience",
            "lịch sử công việc": "experience",
            "kỹ năng": "skills",
            "kỹ năng chuyên môn": "skills",
            "năng lực": "skills",
            "chứng chỉ": "certifications",
            "giấy chứng nhận": "certifications",
            "bằng cấp chuyên môn": "certifications",
            "dự án": "projects",
            "các dự án": "projects",
            "dự án tiêu biểu": "projects",
            "hoạt động": "activities",
            "hoạt động ngoại khóa": "activities",
            "hoạt động xã hội": "activities",
            "giải thưởng": "awards",
            "thành tích": "awards",
            "danh hiệu": "awards",
            "sở thích": "interests",
            "sở trường": "interests",
            "thông tin thêm": "additional",
            "người tham chiếu": "references",
            "tham chiếu": "references",
            # English mappings
            "personal information": "personal_info",
            "contact information": "contact",
            "contact": "contact",
            "objective": "objective",
            "career objective": "objective",
            "summary": "summary",
            "professional summary": "summary",
            "education": "education",
            "academic background": "education",
            "qualifications": "education",
            "experience": "experience",
            "work experience": "experience",
            "employment history": "experience",
            "professional experience": "experience",
            "skills": "skills",
            "technical skills": "skills",
            "core competencies": "skills",
            "competencies": "skills",
            "certifications": "certifications",
            "certificates": "certifications",
            "licenses": "certifications",
            "projects": "projects",
            "key projects": "projects",
            "notable projects": "projects",
            "activities": "activities",
            "extracurricular activities": "activities",
            "volunteer work": "activities",
            "awards": "awards",
            "achievements": "awards",
            "honors": "awards",
            "interests": "interests",
            "hobbies": "interests",
            "additional information": "additional",
            "references": "references",
            "referees": "references",
        }

        return header_mapping.get(header, header.replace(" ", "_"))

    def _build_smart_truncated_content(
        self,
        sections: Dict[str, str],
        priority_sections: List[str],
        max_length: int
    ) -> str:
        """
        Build truncated CV content intelligently based on section priority.

        Ensures high-priority sections (especially experience) are included fully,
        preventing information loss from naive truncation.

        Args:
            sections: Dictionary of section_name -> content
            priority_sections: List of section names in priority order
            max_length: Maximum total character length

        Returns:
            Smartly truncated CV content with sections marked
        """
        smart_content = ""
        remaining_budget = max_length

        logger.info(f"🔧 Building smart content. Total sections: {
                    len(sections)}, Budget: {max_length}")

        # Phase 1: Add priority sections first (guarantee full content)
        for section_name in priority_sections:
            if section_name in sections and remaining_budget > 0:
                section_content = sections[section_name]
                section_text = f"\n\n[{section_name.upper()}]\n{
                    section_content}"

                if len(section_text) <= remaining_budget:
                    smart_content += section_text
                    remaining_budget -= len(section_text)
                    logger.info(f"✅ Added priority section '{section_name}': {
                                len(section_text)} chars, remaining: {remaining_budget}")
                else:
                    # Priority section too large - add as much as possible
                    smart_content += section_text[:remaining_budget]
                    logger.warning(f"⚠️  Priority section '{section_name}' truncated: {
                                   remaining_budget}/{len(section_text)} chars")
                    remaining_budget = 0
                    break

        # Phase 2: Add other sections with remaining budget
        for section_name, content in sections.items():
            if section_name not in priority_sections and remaining_budget > 100:  # Keep min 100 chars for other sections
                section_text = f"\n\n[{section_name.upper()}]\n{content}"

                if len(section_text) <= remaining_budget:
                    smart_content += section_text
                    remaining_budget -= len(section_text)
                    logger.info(
                        f"✅ Added non-priority section '{section_name}': {len(section_text)} chars")
                elif remaining_budget > 200:  # Only add partial if we have decent space
                    smart_content += section_text[:remaining_budget]
                    logger.info(
                        f"⚠️  Non-priority section '{section_name}' partially added: {remaining_budget} chars")
                    remaining_budget = 0
                    break

        logger.info(f"🎯 Smart content built: {
                    len(smart_content)}/{max_length} chars used")
        return smart_content.strip()

    async def _call_ollama(self, prompt: str) -> str:
        """
        Call Ollama API with a prompt and return the response.

        Uses dynamic timeout based on number of pending requests.
        Since Ollama processes requests sequentially, each request needs
        to wait for all pending requests to complete first.

        Timeout calculation: (pending_requests + 1) * BASE_TIMEOUT
        - 1 CV: 120s timeout
        - 2 CVs: 240s timeout (wait for 1st + process 2nd)
        - 5 CVs: 600s timeout (wait for 4 + process 5th)
        """
        global _pending_requests

        # Register this request and calculate timeout
        async with _request_lock:
            _pending_requests += 1
            queue_position = _pending_requests
            dynamic_timeout = queue_position * _BASE_TIMEOUT
            logger.info(f"Ollama request queued. Position: {
                        queue_position}, Timeout: {dynamic_timeout}s")

        try:
            async with httpx.AsyncClient(timeout=dynamic_timeout) as client:
                response = await client.post(
                    f"{self.ollama_url}/api/generate",
                    json={
                        "model": self.model,
                        "prompt": prompt,
                        "stream": False,
                        "options": {
                            "seed": 42,  # Fixed seed for reproducibility
                            # Near-deterministic (0.1 recommended for consistent yet quality results)
                            "temperature": 0.1,
                            "top_p": 0.9,  # Nucleus sampling
                            "num_predict": 2048,  # Max tokens to generate
                        }
                    }
                )
                response.raise_for_status()
                result = response.json()
                logger.info(f"Ollama analysis completed. Queue position was: {
                            queue_position}")
                return result.get("response", "")
        except httpx.TimeoutException:
            logger.error(f"Ollama request timed out after {dynamic_timeout}s")
            raise TimeoutError("AI service request timed out")
        except httpx.HTTPStatusError as e:
            logger.error(f"Ollama HTTP error: {e.response.status_code}")
            raise
        except Exception as e:
            logger.error(f"Ollama request failed: {str(e)}")
            raise
        finally:
            # Always decrement counter when done
            async with _request_lock:
                _pending_requests -= 1
                logger.info(f"Ollama request finished. Remaining in queue: {
                            _pending_requests}")

    async def _update_analysis_status(
        self,
        db: AsyncSession,
        cv_id: uuid.UUID,
        status: models.AnalysisStatus
    ) -> None:
        """
        Update the analysis status in the database.
        """
        stmt = (
            update(models.CVAnalysis)
            .where(models.CVAnalysis.cv_id == cv_id)
            .values(status=status.value)
        )
        await db.execute(stmt)
        await db.commit()

    async def _save_analysis_results(
        self,
        db: AsyncSession,
        cv_id: uuid.UUID,
        results: Dict[str, Any]
    ) -> None:
        """
        Save the AI analysis results to the database.
        """
        # Combine all feedback into ai_feedback JSON field
        feedback = {
            "criteria": results.get("criteria", {}),
            "experience_breakdown": results.get("experience_breakdown", {}),
            "strengths": results.get("strengths", []),
            "improvements": results.get("improvements", []),
            "formatting_feedback": results.get("formatting_feedback", []),
            "ats_hints": results.get("ats_hints", []),
        }

        logger.info(f"💾 SAVING ANALYSIS RESULTS - CV ID: {cv_id}")
        logger.info(f"💾 Overall Score: {results.get('score')}")
        logger.info(f"💾 Experience Score: {results.get(
            'criteria', {}).get('experience', 'N/A')}")
        logger.info(f"💾 Experience Years: {results.get(
            'experience_breakdown', {}).get('total_years', 'N/A')}")

        stmt = (
            update(models.CVAnalysis)
            .where(models.CVAnalysis.cv_id == cv_id)
            .values(
                ai_score=results.get("score"),
                ai_summary=results.get("summary"),
                ai_feedback=feedback,
                extracted_skills=results.get("skills"),
                # Story 5.4: Persist hybrid skill scoring results
                skill_breakdown=results.get("skill_breakdown"),
                skill_categories=results.get("skill_categories"),
                skill_recommendations=results.get("skill_recommendations"),
            )
        )
        await db.execute(stmt)
        await db.commit()

        logger.info(f"✅ ANALYSIS SAVED - CV ID: {cv_id}")


# Global AI service instance
ai_service = AIService()
